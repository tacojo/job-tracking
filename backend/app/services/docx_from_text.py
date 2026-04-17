"""Create a DOCX file from plain text (e.g. tailored CV or cover letter)."""

from io import BytesIO

from docx import Document


def create_docx_from_text(text: str) -> bytes:
    """Create a single-paragraph DOCX from plain text. Returns file bytes."""
    doc = Document()
    if text and text.strip():
        # Split by double newlines for paragraphs, single newlines become line breaks
        for block in text.strip().split("\n\n"):
            p = doc.add_paragraph()
            for line in block.split("\n"):
                if p.runs or p.text:
                    p.add_run("\n")
                p.add_run(line)
    else:
        doc.add_paragraph("(No content)")
    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()
