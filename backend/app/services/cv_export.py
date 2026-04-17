"""Export CV data to DOCX or PDF using templates."""

import io
from pathlib import Path
from typing import Any

from jinja2 import BaseLoader, Environment, select_autoescape

TEMPLATES_DIR = Path(__file__).parent.parent / "templates" / "cv"

# Tailored CV DOCX template names (profile header + body text)
TAILORED_CV_TEMPLATES = ["default", "modern", "minimal"]


def cv_experiences_to_text(
    profile: dict[str, Any], experiences: list[dict[str, Any]]
) -> str:
    """Serialize profile + experiences to plain text for AI tailoring (no formatting)."""
    parts = []
    full_name = (profile.get("full_name") or "").strip()
    tagline = (profile.get("tagline") or "").strip()
    summary = (profile.get("summary") or "").strip()
    if full_name:
        parts.append(full_name)
    if tagline:
        parts.append(tagline)
    if summary:
        parts.append(summary)
    if parts:
        parts.append("")
    parts.append("Experience")
    parts.append("")
    for exp in experiences:
        role = exp.get("role", "") or ""
        employer = exp.get("employer", "") or ""
        parts.append(f"{role} at {employer}")
        start = exp.get("start_date", "") or ""
        end = exp.get("end_date", "") or ""
        loc = (exp.get("location") or "").strip()
        meta = f"{start} - {end}"
        if loc:
            meta = f"{meta}, {loc}"
        parts.append(meta)
        for d in exp.get("details") or []:
            parts.append(f"  • {d}")
        skills = exp.get("skills") or []
        if skills:
            parts.append("Skills: " + ", ".join(skills))
        parts.append("")
    return "\n".join(parts).strip() or "(No CV content)"


def _get_context(
    profile: dict[str, Any], experiences: list[dict[str, Any]]
) -> dict[str, Any]:
    """Build template context from profile and experiences."""
    return {
        "profile": profile,
        "experiences": experiences,
        "full_name": profile.get("full_name", ""),
        "tagline": profile.get("tagline", ""),
        "summary": profile.get("summary", ""),
    }


def export_docx(
    profile: dict[str, Any],
    experiences: list[dict[str, Any]],
    template_name: str = "default",
) -> bytes:
    """Export CV to DOCX using python-docx (programmatic build)."""
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt

    doc = Document()
    ctx = _get_context(profile, experiences)

    # Header
    p = doc.add_paragraph()
    r = p.add_run(ctx["full_name"] or "Your Name")
    r.bold = True
    r.font.size = Pt(18)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    if ctx["tagline"]:
        p = doc.add_paragraph(ctx["tagline"])
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()
    if ctx["summary"]:
        doc.add_paragraph(ctx["summary"])
        doc.add_paragraph()

    # Experience
    p_exp = doc.add_paragraph()
    p_exp.add_run("Experience").bold = True
    doc.add_paragraph()

    for exp in experiences:
        role = exp.get("role", "")
        employer = exp.get("employer", "")
        p_role = doc.add_paragraph()
        p_role.add_run(f"{role} at {employer}").bold = True
        dates = f"{exp.get('start_date', '')} - {exp.get('end_date', '')}"
        loc = exp.get("location", "")
        if loc:
            dates = f"{dates}, {loc}"
        doc.add_paragraph(dates)
        for d in exp.get("details", []):
            doc.add_paragraph(d, style="List Bullet")
        skills = exp.get("skills", [])
        if skills:
            doc.add_paragraph("Skills: " + ", ".join(skills))
        doc.add_paragraph()

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def export_pdf(
    profile: dict[str, Any],
    experiences: list[dict[str, Any]],
    template_name: str = "default",
) -> bytes:
    """Export CV to PDF using HTML template + WeasyPrint."""
    try:
        from weasyprint import HTML
    except ImportError as e:
        raise RuntimeError(
            "WeasyPrint is not installed or its dependencies (Cairo, Pango) are missing. "
            "On Windows, install GTK3-Runtime. See https://doc.courtbouillon.org/weasyprint/"
        ) from e

    tpl_path = TEMPLATES_DIR / "html" / f"{template_name}.html"
    if not tpl_path.exists():
        tpl_path = TEMPLATES_DIR / "html" / "default.html"
    if not tpl_path.exists():
        # Use built-in default HTML
        html_content = _default_html_template()
    else:
        html_content = tpl_path.read_text(encoding="utf-8")

    env = Environment(loader=BaseLoader(), autoescape=select_autoescape())
    template = env.from_string(html_content)
    ctx = _get_context(profile, experiences)
    rendered = template.render(**ctx)

    html = HTML(string=rendered, base_url=str(TEMPLATES_DIR / "html"))
    return html.write_pdf()


def _default_html_template() -> str:
    """Built-in HTML template for PDF export."""
    return """<!DOCTYPE html>
<html>
<head>
<meta charset="utf-8">
<style>
body { font-family: Georgia, serif; max-width: 700px; margin: 40px auto; padding: 20px; line-height: 1.5; }
h1 { font-size: 24px; margin-bottom: 4px; }
.subtitle { color: #444; font-size: 14px; margin-bottom: 20px; }
h2 { font-size: 16px; border-bottom: 1px solid #ccc; margin-top: 24px; }
.exp { margin-bottom: 20px; }
.exp-title { font-weight: bold; font-size: 14px; }
.exp-meta { color: #555; font-size: 13px; margin: 4px 0; }
.exp-details { margin: 8px 0 0 16px; }
.exp-details li { margin: 4px 0; }
.skills { font-size: 12px; color: #666; margin-top: 8px; }
</style>
</head>
<body>
<h1>{{ full_name }}</h1>
<p class="subtitle">{{ tagline }}</p>
{% if summary %}
<p>{{ summary }}</p>
{% endif %}

<h2>Experience</h2>
{% for exp in experiences %}
<div class="exp">
  <div class="exp-title">{{ exp.role }} at {{ exp.employer }}</div>
  <div class="exp-meta">{{ exp.start_date }} - {{ exp.end_date }}{% if exp.location %}, {{ exp.location }}{% endif %}</div>
  <ul class="exp-details">
  {% for d in exp.details %}
    <li>{{ d }}</li>
  {% endfor %}
  </ul>
  {% if exp.skills %}
  <div class="skills">Skills: {{ exp.skills | join(", ") }}</div>
  {% endif %}
</div>
{% endfor %}
</body>
</html>
"""


def export_tailored_cv_docx(
    profile: dict[str, Any],
    tailored_body_text: str,
    template_name: str = "default",
) -> bytes:
    """Build DOCX from profile (header: name, tagline) + tailored body text. Preserves contact details."""
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt

    doc = Document()
    full_name = (profile.get("full_name") or "").strip() or "Your Name"
    tagline = (profile.get("tagline") or "").strip()

    # Header (contact details)
    p = doc.add_paragraph()
    r = p.add_run(full_name)
    r.bold = True
    r.font.size = Pt(18)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if tagline:
        p2 = doc.add_paragraph(tagline)
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()

    # Body: tailored text as paragraphs
    body = (tailored_body_text or "").strip()
    if body:
        for block in body.split("\n\n"):
            block = block.strip()
            if not block:
                continue
            p_para = doc.add_paragraph()
            for i, line in enumerate(block.split("\n")):
                if i > 0:
                    p_para.add_run("\n")
                p_para.add_run(line)
    else:
        doc.add_paragraph("(No content)")

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def list_tailored_cv_templates() -> list[str]:
    """Template names for tailored CV DOCX."""
    return list(TAILORED_CV_TEMPLATES)


def list_templates(format: str) -> list[str]:
    """List available template names for a format (docx or html)."""
    if format == "docx":
        return ["default"]  # Programmatic only for now
    subdir = TEMPLATES_DIR / "html"
    if not subdir.exists():
        return ["default"]
    names = []
    for f in subdir.iterdir():
        if f.suffix == ".html":
            names.append(f.stem)
    return sorted(names) if names else ["default"]
